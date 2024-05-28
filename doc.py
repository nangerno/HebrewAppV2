import docx
from spire.doc import *
from spire.doc.common import *
import shutil

shutil.copy('template.docx', 'new.docx')
document = docx.Document("new.docx")
print("======")
table = document.tables[0]
table.cell(1,1).text = "fffdfssdfs"
document.save("complete.docx")

doc = Document()
doc.LoadFromFile("complete.docx")
doc.SaveToFile("complete.pdf", FileFormat.PDF)
doc.Close()

